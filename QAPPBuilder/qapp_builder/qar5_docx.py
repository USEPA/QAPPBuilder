# qar5_docx.py (qapp_builder)
# !/usr/bin/env python3
# coding=utf-8
# young.daniel@epa.gov

"""Definition of qar5 docx export methods."""


import tempfile
from zipfile import ZipFile
from io import BytesIO
from os import path
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from django.http import HttpResponse, HttpResponseRedirect
from django.templatetags.static import static
from django.utils.text import slugify
from constants.qar5_sectionb import SECTION_B_INFO
from qapp_builder.settings import DEBUG, STATIC_ROOT
from qapp_builder.views import get_qar5_for_user, get_qar5_for_team, \
    get_qapp_info


def export_doc(request, *args, **kwargs):
    """Export a multiple QAPP objects as Word Docx files."""
    if 'user' in request.path:
        user_id = kwargs.get('pk', None)
        team_id = qapp_id = None
    elif 'team' in request.path:
        team_id = kwargs.get('pk', None)
        user_id = qapp_id = None
    else:
        qapp_id = kwargs.get('pk', None)
        team_id = user_id = None

    if qapp_id is None or user_id or team_id:
        if user_id:
            qapp_ids = get_qar5_for_user(
                user_id).values_list('id', flat=True)
        else:
            qapp_ids = get_qar5_for_team(
                team_id).values_list('id', flat=True)

        zip_mem = BytesIO()
        archive = ZipFile(zip_mem, 'w')
        for q_id in qapp_ids:
            resp = export_doc_single(request, pk=q_id)
            filename = resp['filename']
            if filename:
                temp_file_name = '%d_%s' % (q_id, filename)
                with tempfile.SpooledTemporaryFile():
                    archive.writestr(temp_file_name, resp.content)

        archive.close()
        response = HttpResponse(
            zip_mem.getvalue(), content_type='application/force-download')
        response['Content-Disposition'] = \
            'attachment; filename="%s_qapps.zip"' % request.user.username
        response['Content-length'] = zip_mem.tell()
        return response


def add_custom_heading(document, text, level):
    """Add centered headers to a docx file."""
    heading_style = 'custom_header_%d' % level
    paragraph = document.add_paragraph(text, heading_style)
    return paragraph


def add_center_heading(document, text, level):
    """Add centered headers to a docx file."""
    paragraph = add_custom_heading(document, text, level)
    # pylint: disable=no-member
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def set_table_row_height(table):
    """Set minimum row height and alignment for a table."""
    for row in table.rows:
        # pylint: disable=no-member
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.35)
        for cell in row.cells:
            # pylint: disable=no-member
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_custom_headers(document):
    """
    Create and inject custom headers into a docx.

    This functionality is necessary so the cover page isn't added
    to the auto-generated Table of Contents.
    """
    custom_header_1 = document.styles.add_style(
        # pylint: disable=no-member
        'custom_header_1', WD_STYLE_TYPE.PARAGRAPH)
    custom_header_1.font.size = Pt(14)
    custom_header_1.font.bold = True
    # pylint: disable=no-member
    custom_header_1.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1

    custom_header_2 = document.styles.add_style(
        # pylint: disable=no-member
        'custom_header_2', WD_STYLE_TYPE.PARAGRAPH)
    custom_header_2.font.size = Pt(13)
    custom_header_2.font.bold = True
    # pylint: disable=no-member
    custom_header_2.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1

    custom_header_3 = document.styles.add_style(
        # pylint: disable=no-member
        'custom_header_3', WD_STYLE_TYPE.PARAGRAPH)
    custom_header_3.font.size = Pt(11)
    custom_header_3.font.bold = True
    # pylint: disable=no-member
    custom_header_3.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1


def create_toc(document):
    """Set up the Table of Contents page."""
    add_custom_heading(document, 'A.2 Table of Contents', level=2)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    # creates a new element
    field_char = OxmlElement('w:fldChar')
    # sets attribute on element
    field_char.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    # sets attribute on element
    instr_text.set(qn('xml:space'), 'preserve')
    # change 1-3 depending on heading levels you need
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    field_char2 = OxmlElement('w:fldChar')
    field_char2.set(qn('w:fldCharType'), 'separate')
    field_char3 = OxmlElement('w:t')
    field_char3.text = "Right-click to update Table of Contents."
    field_char2.append(field_char3)

    field_char4 = OxmlElement('w:fldChar')
    field_char4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(field_char)
    r_element.append(instr_text)
    r_element.append(field_char2)
    r_element.append(field_char4)


# py-lint: disable=no-member
def export_doc_single(request, *args, **kwargs):
    """Export a single QAPP object as a Word Docx file."""
    qapp_id = kwargs.get('pk', None)
    qapp_info = get_qapp_info(request.user, qapp_id)
    qapp_info['qapp'] = qapp_info['qapp']

    if not qapp_info:
        return HttpResponseRedirect(request)

    filename = '%s.docx' % slugify(qapp_info['qapp'].title)

    document = Document()
    add_custom_headers(document)
    styles = document.styles

    # #################################################
    # BEGIN COVER PAGE
    # #################################################
    # Coversheet with signatures section:
    run = document.add_paragraph().add_run()

    try:
        if DEBUG:
            logo = path.join(STATIC_ROOT, 'EPA_Files', 'logo.png')
            qual_assur_proj_plan = path.join(
                STATIC_ROOT, 'images', 'quality_assurance_project_plan.PNG')
        else:
            logo = static('logo.png')
            qual_assur_proj_plan = static('quality_assurance_project_plan.PNG')

        run.add_picture(logo, width=Inches(1.5))
        run.add_text('\t\t\t')
        run.add_picture(qual_assur_proj_plan, width=Inches(3))

    except FileNotFoundError:
        print('couldn\'t find the static images!')

    section_incomplete_message = 'This section has errors or is ' + \
        'incomplete. Please go back to the web tool and correct ' + \
        'this section, then try exporting it again.'

    # The rest of the document will be WD_ALIGN_PARAGRAPH.CENTER

    # blank line
    add_center_heading(document, 'Office of Research and Development', 1)
    add_center_heading(document, qapp_info['qapp'].division.name, 1)
    # blank line
    # Next few sections are from the qapp object
    add_center_heading(document, qapp_info['qapp'].division_branch, level=3)
    # blank line
    add_center_heading(document, 'EPA Project Lead', level=2)
    for lead in qapp_info['qapp_leads']:
        add_center_heading(document, lead.name, level=3)
    # blank line
    add_center_heading(document, qapp_info['qapp'].intra_extra, level=3)
    add_center_heading(document, qapp_info['qapp'].qa_category, level=3)
    add_center_heading(document, qapp_info['qapp'].revision_number, level=3)
    add_center_heading(document, str(qapp_info['qapp'].date), level=3)
    # blank line
    add_center_heading(document, 'Prepared By', level=2)
    add_center_heading(
        document, '%s %s' % (qapp_info['qapp'].prepared_by.first_name,
                             qapp_info['qapp'].prepared_by.last_name,),
        level=3)
    # blank line
    add_center_heading(document, qapp_info['qapp'].strap, level=3)
    add_center_heading(document, qapp_info['qapp'].tracking_id, level=3)

    # #################################################
    # END COVER PAGE
    # BEGIN APPROVAL PAGE
    # #################################################

    add_custom_heading(document, 'A.1 Approval Page', level=2)
    # Signature grid ...
    num_signatures = 0
    num_signatures = len(qapp_info.get('signatures', []))
    table = document.add_table(rows=6+num_signatures, cols=12)
    table.style = styles['Table Grid']

    set_table_row_height(table)

    if qapp_info.get('qapp_approval', False):
        row_cells = table.rows[0].cells
        row_cells[0].text = 'QA Project Plan Title:'
        row_cells[0].merge(row_cells[3])
        row_cells[4].text = qapp_info['qapp_approval'].project_plan_title
        row_cells[4].merge(row_cells[11])

        row_cells = table.rows[1].cells
        row_cells[0].text = 'QA Activity Number:'
        row_cells[0].merge(row_cells[3])
        row_cells[4].text = qapp_info['qapp_approval'].activity_number
        row_cells[4].merge(row_cells[11])

    row_cells = table.rows[2].cells
    row_cells[0].text = 'If Intramural or Extramural, EPA Project Approvals'
    row_cells[0].merge(row_cells[11])

    iter_count = 0
    # Start with row 3 + iter_count++
    for sig in qapp_info.get('signatures', []):
        if not sig.contractor:
            row_cells = table.rows[3 + iter_count].cells
            row_cells[0].text = 'Name:'
            row_cells[1].text = sig.name
            row_cells[1].merge(row_cells[3])

            row_cells[4].text = 'Signature/Date:'
            row_cells[4].merge(row_cells[5])
            row_cells[6].merge(row_cells[11])
            iter_count += 1

    # Always insert a blank entry for hand-written approval sigs
    row_cells = table.rows[3 + iter_count].cells
    row_cells[0].text = 'Name:'
    row_cells[1].merge(row_cells[3])
    row_cells[4].text = 'Signature/Date:'
    row_cells[4].merge(row_cells[5])
    row_cells[6].merge(row_cells[11])

    row_cells = table.rows[4 + iter_count].cells
    row_cells[0].text = 'If Extramural, Contractor Project Approvals'
    row_cells[0].merge(row_cells[11])

    for sig in qapp_info.get('signatures', []):
        if sig.contractor:
            row_cells = table.rows[5 + iter_count].cells
            row_cells[0].text = 'Name:'
            row_cells[1].text = sig.name
            row_cells[1].merge(row_cells[3])

            row_cells[4].text = 'Signature/Date:'
            row_cells[4].merge(row_cells[5])
            row_cells[6].merge(row_cells[11])
            iter_count += 1

    # Always insert a blank entry for hand-written approval sigs
    row_cells = table.rows[5 + iter_count].cells
    row_cells[0].text = 'Name:'
    row_cells[1].merge(row_cells[3])
    row_cells[4].text = 'Signature/Date:'
    row_cells[4].merge(row_cells[5])
    row_cells[6].merge(row_cells[11])
    document.add_page_break()

    # #################################################
    # END APPROVAL PAGE
    # BEGIN ToC PAGE
    # #################################################

    create_toc(document)
    document.add_heading('Definitions and Acronyms', level=2)

    if qapp_info['section_a']:
        if qapp_info['section_a'].a2:
            document.add_paragraph(
                qapp_info['section_a'].a2,
                styles['No Spacing'])

        if qapp_info['section_a'].a2_keywords:
            document.add_heading('Keywords', level=2)
            document.add_paragraph(
                qapp_info['section_a'].a2_keywords,
                styles['No Spacing'])

    document.add_page_break()

    # #################################################
    # END ToC PAGE
    # BEGIN Everything Else PAGE
    # #################################################

    #  1) Heading 1 - Revision History
    document.add_heading('Revision History', level=1)
    #  2) Table Label
    document.add_heading('Table 1 QAPP Revision History', level=3)
    #  3) Table (revision history)
    num_revisions = len(qapp_info['revisions'])
    table = document.add_table(rows=1+num_revisions, cols=3)
    table.style = styles['Light List']
    row_cells = table.rows[0].cells
    row_cells[0].text = 'Revision Number'
    row_cells[1].text = 'Date Approved'
    row_cells[2].text = 'Revision'

    iter_count = 0
    for rev in qapp_info['revisions']:
        row_cells = table.rows[1+iter_count].cells
        row_cells[0].text = rev.revision
        row_cells[1].text = str(rev.effective_date)
        row_cells[2].text = rev.description
        iter_count += 1

    # Section A
    document.add_heading('Section A - Executive Summary', level=1)
    if qapp_info['section_a']:
        document.add_heading('A.3 Distribution List', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a3,
            styles['No Spacing'])
        document.add_heading('A.4 Project Task Organization', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a4,
            styles['No Spacing'])
        document.add_heading('A.5 Problem Definition Background', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a5,
            styles['No Spacing'])
        document.add_heading('A.6 Project Description', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a6,
            styles['No Spacing'])
        document.add_heading('A.7 Quality Objectives and Criteria', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a7,
            styles['No Spacing'])
        document.add_heading('A.8 Special Training Certification', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a8,
            styles['No Spacing'])
        document.add_heading('A.9 Documents and Records', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a9,
            styles['No Spacing'])
    else:
        document.add_paragraph(section_incomplete_message,
                               styles['No Spacing'])

    # Section B
    if qapp_info['section_b']:
        for sectionb in qapp_info['section_b']:
            sectionb_type = sectionb.sectionb_type.name
            document.add_heading('Section B - %s' % sectionb_type, level=1)
            section_b_info = SECTION_B_INFO[sectionb_type]
            for key in section_b_info:
                val = getattr(sectionb, key, '')
                if section_b_info[key].get('heading', False):
                    document.add_heading(section_b_info[key]['heading'],
                                         level=2)
                document.add_heading(section_b_info[key]['label'], level=3)
                document.add_paragraph(val, styles['No Spacing'])
    else:
        document.add_heading('Section B', level=1)
        document.add_paragraph(section_incomplete_message,
                               styles['No Spacing'])

    # Section C
    document.add_heading('Section C', level=1)
    if qapp_info['section_c']:
        document.add_heading(
            'C.1 Assessments and Response Actions', level=2)
        document.add_paragraph(
            qapp_info['section_c'].c1,
            styles['No Spacing'])
        document.add_heading('C.2 Reports to Management', level=2)
        document.add_paragraph(
            qapp_info['section_c'].c2,
            styles['No Spacing'])
        # document.add_heading('C.3 Quality Metrics (QA/QC Checks)', level=2)
        # document.add_paragraph(
        #     qapp_info['section_c'].c3,
        #     styles['No Spacing'])
    else:
        document.add_heading(
            'C.1 Assessments and Response Actions', level=2)
        document.add_heading('C.2 Reports to Management', level=2)

    # Section D
    if qapp_info['section_d']:
        document.add_heading('Section D', level=1)
        document.add_heading(
            'D.1 Data Review, Verification, and Validation', level=2)
        document.add_paragraph(
            qapp_info['section_d'].d1,
            styles['No Spacing'])
        document.add_heading(
            'D.2 Verification and Validation Methods', level=2)
        document.add_paragraph(
            qapp_info['section_d'].d2,
            styles['No Spacing'])
        document.add_heading(
            'D.3 Reconciliation with User Requirements', level=2)
        document.add_paragraph(
            qapp_info['section_d'].d3,
            styles['No Spacing'])
    else:
        document.add_paragraph(section_incomplete_message,
                               styles['No Spacing'])

    # References
    document.add_heading('References', level=1)
    if qapp_info['references']:
        run = document.add_paragraph().add_run()
        run.add_text(qapp_info['references'].references)
        # document.add_paragraph(
        #     qapp_info['references'].references.replace('\r\n\r\n', '\r\n'),
        #     styles['No Spacing'])
    else:
        document.add_paragraph(section_incomplete_message,
                               styles['No Spacing'])

    content_type = 'application/vnd.openxmlformats-officedocument.' + \
                   'wordprocessingml.document'
    response = HttpResponse(content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % filename
    document.save(response)
    response['filename'] = filename
    return response
