# docx.py (QAPP_Builder)
# !/usr/bin/env python3
# coding=utf-8
# young.daniel@epa.gov


"""Working with .docx module to edit QAPPs built in app."""

# NOTE: import docx NOT python-docx
import docx
# create an instance of a word document
# this should be add existing QAPP created by user
# should bring user to any QAPP they have created, listed on edit QAPP section

doc = docx.Document()
# add a heading of level 0 (largest heading)
doc.add_heading('Heading for the document', 0)
# add a paragraph and store
# the object in a variable
doc_para = doc.add_paragraph('Your paragraph goes here, ')
# add a run i.e, style like
# bold, italic, underline, etc.
doc_para.add_run('hey there, bold here').bold = True
doc_para.add_run(', and ')
doc_para.add_run('these words are italic').italic = True
# add a page break to start a new page
doc.add_page_break()
# add a heading of level 2
doc.add_heading('Heading level 2', 2)
# pictures can also be added to our word document
# width is optional
doc.add_picture('path_to_picture')
# now save the document to a location
doc.save('path_to_document')
